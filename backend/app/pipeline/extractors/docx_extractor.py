"""Экстрактор DOCX (Word). Кейс: Клиника 1, прайс 2024.

Особенности формата (раздел 7.3):
  * Прайс — это ОДНА таблица на 2727 строк (Код / Наименование / Стоимость).
    Поэтому идём построчно, потоково, не собирая гигантскую структуру в памяти.
  * В документе могут жить непринятые правки (track changes). Финальный текст —
    это документ с ПРИНЯТЫМИ изменениями: <w:del> выкидываем целиком,
    <w:ins> разворачиваем (оставляем вложенные ранчики w:r). Делаем это на
    уровне XML до открытия python-docx, потому что python-docx сам правки не
    принимает и может вернуть зачёркнутый/старый текст.
  * Разделы («Раздел 1.Консультации специалистов») — это строки, где все ячейки
    объединены (один и тот же текст) и нет цены. Они становятся category для
    последующих позиций.

Экстрактор заполняет только service_name_raw / service_code_source / prices /
category / source_row. Раскладка тарифов в резидент/нерезидент — позже, в
price_parser. Карта prices: «подпись колонки -> Decimal».
"""

from __future__ import annotations

import os
import re
import tempfile
import zipfile

import docx

from app.pipeline.base import ExtractedItem, ExtractionResult, Extractor
from app.pipeline.price_parser import parse_amount

# Якорные слова заголовка таблицы (дублируем локально, чтобы не тащить геометрию
# PDF; для docx ячейки уже разложены по колонкам самим Word-ом).
_NAME_HINTS = ("наименование", "услуга", "услуги", "название", "обследование")
_CODE_HINTS = ("код", "шифр", "тарификатор", "№", "No")
_PRICE_HINTS = (
    "цена", "стоимость", "тариф", "сумма", "тенге", "тг",
    "резидент", "нерезидент", "гражд", "иностран", "снг", "зарубеж",
    "первичный", "повторный", "страхов", "руб",
)
_HEADER_ANCHORS = _NAME_HINTS + _CODE_HINTS + _PRICE_HINTS

# Строки-итоги, которые пропускаем (раздел 7.7, правило).
_TOTAL_RE = re.compile(r"\b(итого|всего|подытог|сумма по разделу)\b", re.IGNORECASE)

# Ограничение на размер сырого текста для аудита.
_RAW_LIMIT = 500_000

# Пространства имён WordprocessingML.
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Регэкспы по байтам XML: убрать <w:del>…</w:del> целиком и развернуть <w:ins>.
# Работаем с любым префиксом (w:del, обычно так), допускаем самозакрытие и атрибуты.
_DEL_RE = re.compile(rb"<w:del\b[^>]*?/>|<w:del\b[^>]*?>.*?</w:del>", re.DOTALL)
_INS_OPEN_RE = re.compile(rb"<w:ins\b[^>]*?>")
_INS_CLOSE_RE = re.compile(rb"</w:ins>")
_INS_SELF_RE = re.compile(rb"<w:ins\b[^>]*?/>")


def accept_tracked_changes(path: str) -> str:
    """Принимает все правки (track changes) на уровне XML и возвращает путь к
    очищенной временной копии .docx. При любой ошибке возвращает исходный путь
    (вызывающая сторона уже залогирует через warn)."""
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx", prefix="docx_accept_")
    os.close(tmp_fd)
    try:
        with zipfile.ZipFile(path, "r") as zin:
            names = zin.namelist()
            # Чистим основной документ и связанные части (headers/footers тоже
            # могут содержать правки, но прайс — в document.xml; почистим всё,
            # где встречаются track-change теги).
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in names:
                    data = zin.read(item)
                    if item.endswith(".xml") and (b"<w:del" in data or b"<w:ins" in data):
                        data = _accept_xml_bytes(data)
                    zout.writestr(item, data)
        return tmp_path
    except Exception:
        # Не получилось — убираем мусорный временный файл и сигналим наверх,
        # вернув исходный путь.
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
        return path


def _accept_xml_bytes(data: bytes) -> bytes:
    """Удаляет <w:del>…</w:del> и разворачивает <w:ins> в байтовом XML."""
    # 1) Удаляем удалённые фрагменты целиком (вместе с w:delText внутри).
    data = _DEL_RE.sub(b"", data)
    # 2) Разворачиваем вставки: убираем сами теги w:ins, сохраняя их содержимое.
    data = _INS_SELF_RE.sub(b"", data)
    data = _INS_OPEN_RE.sub(b"", data)
    data = _INS_CLOSE_RE.sub(b"", data)
    return data


class DocxExtractor(Extractor):
    """Экстрактор для Word-прайсов. Одна большая таблица, потоковый разбор."""

    name = "docx"

    def can_handle(self, file_path: str, file_format: str) -> bool:
        return file_format == "docx"

    def extract(self, file_path: str) -> ExtractionResult:
        result = ExtractionResult(extractor_used=self.name, page_count=0)

        # Шаг 1. Принимаем правки (финальный текст). При ошибке — исходный файл.
        clean_path = accept_tracked_changes(file_path)
        cleaned = clean_path != file_path
        if not cleaned:
            result.warn("Не удалось принять правки track changes — читаю исходный файл")

        try:
            try:
                document = docx.Document(clean_path)
            except Exception as exc:
                # Очищенная копия битая — пробуем исходник.
                if cleaned:
                    result.warn(f"Очищенный docx не открылся ({exc}); читаю оригинал")
                    document = docx.Document(file_path)
                else:
                    raise

            raw_parts: list[str] = []
            self._collect_paragraph_text(document, raw_parts)

            # Шаг 2. Идём по всем таблицам построчно (streaming).
            for table in document.tables:
                self._extract_table(table, result, raw_parts)

            result.raw_content = self._join_raw(raw_parts)
        except Exception as exc:
            # Защита верхнего уровня: не валим pipeline из-за одного файла.
            result.warn(f"Критическая ошибка разбора docx: {exc}")
        finally:
            if cleaned:
                _safe_unlink(clean_path)

        return result

    # --- внутреннее ---------------------------------------------------------

    def _collect_paragraph_text(self, document, raw_parts: list[str]) -> None:
        """Собирает текст параграфов тела (вне таблиц) для raw_content."""
        try:
            for para in document.paragraphs:
                text = para.text.strip()
                if text:
                    raw_parts.append(text)
        except Exception:
            # Параграфы не критичны для извлечения позиций.
            pass

    def _extract_table(self, table, result: ExtractionResult, raw_parts: list[str]) -> None:
        """Потоково разбирает одну таблицу: ищет заголовок, мапит колонки,
        читает позиции построчно и детектит разделы."""
        rows = table.rows
        header_idx, cmap = self._find_header(rows)
        if header_idx < 0 or cmap is None:
            result.warn("В таблице не найден заголовок (Код/Наименование/Стоимость) — пропускаю")
            # Всё равно положим текст таблицы в аудит.
            self._dump_table_text(rows, raw_parts)
            return

        name_idx = cmap["name_idx"]
        code_idx = cmap["code_idx"]
        price_cols = cmap["price_cols"]  # list[(idx, label)]

        current_category: str | None = None

        for r in range(header_idx + 1, len(rows)):
            try:
                cells = self._row_cells(rows[r])
            except Exception as exc:
                result.warn(f"Строка {r}: не прочиталась ({exc})")
                continue
            if not cells:
                continue

            joined = " ".join(c for c in cells if c).strip()
            if joined:
                raw_parts.append(joined)
            if not joined:
                continue

            # Итоги/подытоги — пропускаем.
            if _TOTAL_RE.search(joined):
                continue

            # Раздел: строка с объединёнными ячейками — все непустые ячейки несут
            # один и тот же текст (Word дублирует merged-cell). Проверяем ЭТО
            # раньше цены: иначе число внутри названия раздела («Раздел 1.») будет
            # ошибочно принято за стоимость. У настоящей позиции ячейки различны
            # (код != наименование != цена).
            non_empty = [c for c in cells if c]
            distinct = set(non_empty)
            if non_empty and len(distinct) <= 1:
                current_category = self._clean_category(non_empty[0])
                continue

            prices = self._read_prices(cells, price_cols)

            # Обычная позиция.
            name = self._cell(cells, name_idx).strip()
            code = self._cell(cells, code_idx).strip()
            if not name:
                # Иногда имя пустое, но есть код+цена — это, скорее, мусор.
                if not prices:
                    continue
                # Без названия позиция бесполезна — фиксируем и пропускаем.
                result.warn(f"Строка {r}: пустое наименование — пропускаю")
                continue
            if not prices:
                # Строка без цены и без признаков раздела (разные ячейки) —
                # пропускаем как заголовок-подраздел или мусор, но не теряем в аудите.
                continue

            item = ExtractedItem(
                service_name_raw=name,
                service_code_source=code or None,
                prices=prices,
                category=current_category,
                source_row=r,
            )
            result.items.append(item)

    def _find_header(self, rows, max_scan: int = 25):
        """Сканирует первые строки, ищет заголовок по якорным словам и строит
        карту колонок. Возвращает (header_idx, cmap|None)."""
        best_idx, best_score, best_cells = -1, 0, None
        scan = min(max_scan, len(rows))
        for i in range(scan):
            try:
                cells = self._row_cells(rows[i])
            except Exception:
                continue
            low = " ".join(cells).lower()
            score = sum(1 for a in _HEADER_ANCHORS if a in low)
            if score > best_score:
                best_score, best_idx, best_cells = score, i, cells
        if best_score < 2 or best_cells is None:
            return -1, None
        return best_idx, self._map_columns(best_cells)

    def _map_columns(self, header_cells: list[str]) -> dict:
        """По ячейкам заголовка строит карту: name_idx, code_idx, price_cols."""
        name_idx: int | None = None
        code_idx: int | None = None
        price_cols: list[tuple[int, str]] = []
        for idx, cell in enumerate(header_cells):
            low = cell.lower().strip()
            if not low:
                continue
            if name_idx is None and any(h in low for h in _NAME_HINTS):
                name_idx = idx
            elif code_idx is None and any(h in low for h in _CODE_HINTS):
                code_idx = idx
            elif any(h in low for h in _PRICE_HINTS):
                price_cols.append((idx, cell.strip()))
        # Фолбэк: если ценовых колонок нет, считаем последнюю колонку ценой.
        if not price_cols and header_cells:
            last = len(header_cells) - 1
            label = header_cells[last].strip() or "Стоимость"
            if last not in (name_idx, code_idx):
                price_cols.append((last, label))
        return {"name_idx": name_idx, "code_idx": code_idx, "price_cols": price_cols}

    def _read_prices(self, cells: list[str], price_cols: list[tuple[int, str]]) -> dict:
        """Читает цены из ценовых колонок, ключ — подпись колонки. Пропускает
        ячейки, где parse_amount вернул None."""
        prices: dict[str, object] = {}
        for idx, label in price_cols:
            if idx >= len(cells):
                continue
            amount = parse_amount(cells[idx])
            if amount is not None:
                prices[label] = amount
        return prices

    @staticmethod
    def _row_cells(row) -> list[str]:
        """Текст ячеек строки. Объединённые ячейки python-docx отдаёт повторяясь
        — это нам и нужно для детекции разделов (len(set)==1)."""
        return [c.text.strip() for c in row.cells]

    @staticmethod
    def _cell(cells: list[str], idx: int | None) -> str:
        if idx is None or idx >= len(cells):
            return ""
        return cells[idx]

    @staticmethod
    def _clean_category(text: str) -> str:
        """Нормализует подпись раздела: убирает префикс «Раздел N.» и лишние
        пробелы, оставляя содержательное название."""
        cleaned = re.sub(r"^\s*раздел\s*", "", text, flags=re.IGNORECASE)
        cleaned = re.sub(r"^\s*[\d.]+\s*", "", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
        return cleaned or text.strip()

    def _dump_table_text(self, rows, raw_parts: list[str]) -> None:
        """Сваливает текст таблицы в аудит, когда заголовок не распознан."""
        for row in rows:
            try:
                joined = " ".join(c for c in self._row_cells(row) if c).strip()
            except Exception:
                continue
            if joined:
                raw_parts.append(joined)

    @staticmethod
    def _join_raw(raw_parts: list[str]) -> str:
        """Склеивает части в сырой текст, обрезая по лимиту ~500k символов."""
        out: list[str] = []
        size = 0
        for part in raw_parts:
            size += len(part) + 1
            out.append(part)
            if size >= _RAW_LIMIT:
                break
        text = "\n".join(out)
        return text[:_RAW_LIMIT]


def _safe_unlink(path: str) -> None:
    try:
        os.unlink(path)
    except OSError:
        pass
